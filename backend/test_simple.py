#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys
from pathlib import Path

def test_ocr_direct():
    """直接测试OCR功能"""
    print("🔍 直接测试OCR功能")
    print("=" * 50)
    
    try:
        import pytesseract
        from PIL import Image
        
        print(f"pytesseract.tesseract_cmd: {pytesseract.pytesseract.tesseract_cmd}")
        print(f"TESSDATA_PREFIX: {os.environ.get('TESSDATA_PREFIX', '未设置')}")
        
        # 检查tesseract是否可用
        try:
            version = pytesseract.get_tesseract_version()
            print(f"✅ Tesseract版本: {version}")
        except Exception as e:
            print(f"❌ 无法获取Tesseract版本: {e}")
            return False
        
        # 测试OCR
        test_image = Image.new('RGB', (200, 100), color='white')
        text = pytesseract.image_to_string(test_image, lang='eng')
        print(f"✅ OCR测试成功: '{text.strip()}'")
        return True
        
    except Exception as e:
        print(f"❌ OCR测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_word_direct():
    """直接测试Word处理功能"""
    print("\n📄 直接测试Word处理功能")
    print("=" * 50)
    
    try:
        from docx import Document
        
        # 创建一个测试Word文档
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        
        # 添加表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '品名'
        header_cells[1].text = '规格型号'
        header_cells[2].text = '数量'
        
        # 添加数据
        data_cells = table.rows[1].cells
        data_cells[0].text = '阀门'
        data_cells[1].text = 'DN100'
        data_cells[2].text = '5'
        
        # 保存测试文件
        test_file = "test_word.docx"
        doc.save(test_file)
        print(f"✅ 创建测试Word文件: {test_file}")
        
        # 测试读取
        doc2 = Document(test_file)
        print(f"✅ 读取Word文件成功，包含{len(doc2.tables)}个表格")
        
        # 测试表格内容
        if len(doc2.tables) > 0:
            table = doc2.tables[0]
            print(f"✅ 表格大小: {len(table.rows)}行 x {len(table.columns)}列")
            
            # 提取数据
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
                print(f"   行数据: {row_data}")
        
        # 清理
        os.remove(test_file)
        return True
        
    except Exception as e:
        print(f"❌ Word处理测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_parse_function():
    """测试parse_file_to_excel函数"""
    print("\n🔧 测试parse_file_to_excel函数")
    print("=" * 50)
    
    try:
        # 导入函数
        sys.path.append('.')
        from main import parse_file_to_excel
        
        # 创建测试图片
        from PIL import Image
        test_image = Image.new('RGB', (400, 200), color='white')
        test_image.save("test_image.png")
        
        # 创建测试Word文件
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        header_cells = table.rows[0].cells
        header_cells[0].text = '品名'
        header_cells[1].text = '规格型号'
        header_cells[2].text = '数量'
        data_cells = table.rows[1].cells
        data_cells[0].text = '阀门'
        data_cells[1].text = 'DN100'
        data_cells[2].text = '5'
        doc.save("test_word.docx")
        
        # 测试目录
        test_dir = "test_output"
        os.makedirs(test_dir, exist_ok=True)
        
        # 测试图片处理
        print("📤 测试图片处理...")
        with open("test_image.png", 'rb') as f:
            file_bytes = f.read()
            try:
                excel_name = parse_file_to_excel(file_bytes, "test_image.png", test_dir)
                print(f"✅ 图片处理成功: {excel_name}")
            except Exception as e:
                print(f"❌ 图片处理失败: {e}")
        
        # 测试Word处理
        print("📤 测试Word处理...")
        with open("test_word.docx", 'rb') as f:
            file_bytes = f.read()
            try:
                excel_name = parse_file_to_excel(file_bytes, "test_word.docx", test_dir)
                print(f"✅ Word处理成功: {excel_name}")
            except Exception as e:
                print(f"❌ Word处理失败: {e}")
        
        # 清理
        os.remove("test_image.png")
        os.remove("test_word.docx")
        import shutil
        shutil.rmtree(test_dir)
        
        return True
        
    except Exception as e:
        print(f"❌ parse_file_to_excel测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_generate_standard_model_from_multicell():
    import pandas as pd
    from valve_model_generator import parse_valve_info
    # 自动用第3行作为表头
    df = pd.read_excel(r'C:/Users/Administrator/Desktop/quote-system/quote-system/backend/test_input.xlsx', header=2)
    # 打印所有列名，便于调试
    print('实际表头:', df.columns.tolist())
    # 增强备注列检测，支持多种关键词
    note_keywords = ['备注', '注', '★注', '备注栏']
    note_col = None
    for col in df.columns:
        col_no_space = str(col).replace(' ', '')
        for kw in note_keywords:
            if kw in col_no_space:
                note_col = col
                break
        if note_col:
            break
    if note_col is None:
        print('未找到包含“备注/注/★注/备注栏”的列，将不合并备注信息。')
    # 合并品名信息
    def build_full_name(row):
        name = str(row.get('项目名称', ''))
        note = str(row.get(note_col, '')) if note_col else ''
        return f"{name} {note}".strip()
    df['品名'] = df.apply(build_full_name, axis=1)
    df['标准型号'] = df.apply(lambda row: parse_valve_info(row['品名'], ''), axis=1)
    # 打印和保存结果
    cols_to_show = ['项目名称', '品名', '标准型号']
    if note_col:
        cols_to_show.insert(1, note_col)
    print(df[cols_to_show])
    output_path = r'C:/Users/Administrator/Desktop/quote-system/quote-system/backend/test_output_with_model.xlsx'
    df[cols_to_show].to_excel(output_path, index=False)
    print(f"已保存结果到: {output_path}")

def main():
    """主测试函数"""
    print("🚀 开始简化测试")
    print("=" * 60)
    
    # 1. 测试OCR
    ocr_ok = test_ocr_direct()
    
    # 2. 测试Word处理
    word_ok = test_word_direct()
    
    # 3. 测试parse函数
    parse_ok = test_parse_function()
    
    # 总结
    print("\n" + "=" * 60)
    print("📋 测试总结")
    print("=" * 60)
    print(f"OCR功能: {'✅ 正常' if ocr_ok else '❌ 异常'}")
    print(f"Word处理: {'✅ 正常' if word_ok else '❌ 异常'}")
    print(f"parse函数: {'✅ 正常' if parse_ok else '❌ 异常'}")
    
    if not ocr_ok:
        print("\n🔧 OCR问题建议:")
        print("1. 检查 tesseract.exe 路径: D:\\tupian\\tesseract.exe")
        print("2. 设置环境变量: TESSDATA_PREFIX=D:\\tupian\\tessdata")
        print("3. 确保 tesseract.exe 有执行权限")
    
    if not word_ok:
        print("\n🔧 Word处理问题建议:")
        print("1. 检查 python-docx 包: pip install python-docx")
        print("2. 确保Word文件包含标准表格")
    
    if not parse_ok:
        print("\n🔧 parse函数问题建议:")
        print("1. 检查所有依赖包是否正确安装")
        print("2. 查看详细错误信息")

if __name__ == "__main__":
    test_generate_standard_model_from_multicell()
    main() 