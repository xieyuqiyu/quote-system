#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import requests
import os
import sys
from pathlib import Path

def test_ocr_environment():
    """测试OCR环境"""
    print("🔍 测试OCR环境")
    print("=" * 50)
    
    try:
        import pytesseract
        from PIL import Image
        
        # 检查配置
        print(f"pytesseract.tesseract_cmd: {pytesseract.pytesseract.tesseract_cmd}")
        print(f"TESSDATA_PREFIX: {os.environ.get('TESSDATA_PREFIX', '未设置')}")
        
        # 测试OCR
        test_image = Image.new('RGB', (200, 100), color='white')
        text = pytesseract.image_to_string(test_image, lang='eng')
        print(f"✅ OCR测试成功: '{text.strip()}'")
        return True
        
    except Exception as e:
        print(f"❌ OCR环境测试失败: {e}")
        return False

def test_word_processing():
    """测试Word文件处理"""
    print("\n📄 测试Word文件处理")
    print("=" * 50)
    
    try:
        from docx import Document
        
        # 创建一个测试Word文档
        doc = Document()
        table = doc.add_table(rows=3, cols=3)
        
        # 添加表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '品名'
        header_cells[1].text = '规格型号'
        header_cells[2].text = '数量'
        
        # 添加数据
        data_cells = table.rows[1].cells
        data_cells[0].text = '阀门'
        data_cells[1].text = 'DN100'
        data_cells[2].text = '5'
        
        # 保存测试文件
        test_file = "test_word.docx"
        doc.save(test_file)
        print(f"✅ 创建测试Word文件: {test_file}")
        
        # 测试读取
        doc2 = Document(test_file)
        print(f"✅ 读取Word文件成功，包含{len(doc2.tables)}个表格")
        
        # 清理
        os.remove(test_file)
        return True
        
    except Exception as e:
        print(f"❌ Word处理测试失败: {e}")
        return False

def test_file_upload_api():
    """测试文件上传API"""
    print("\n📤 测试文件上传API")
    print("=" * 50)
    
    # 服务器地址
    base_url = "http://localhost:8001"
    
    # 登录测试
    print("🔐 登录测试...")
    try:
        response = requests.post(f"{base_url}/api/login", auth=("admin", "admin123"))
        if response.status_code == 200:
            print("✅ 登录成功")
        else:
            print(f"❌ 登录失败: {response.status_code} - {response.text}")
            return False
    except Exception as e:
        print(f"❌ 登录请求失败: {e}")
        return False
    
    # 测试文件列表
    test_files = []
    
    # 检查存在的文件
    possible_files = [
        "../9ce048b41c461644f43ceca6dc11ad5d.jpg",
        "../123.png", 
        "../3 (2)_01.png",
        "../123.docx"
    ]
    
    for file_path in possible_files:
        if os.path.exists(file_path):
            test_files.append(file_path)
            print(f"✅ 找到测试文件: {file_path}")
        else:
            print(f"⚠️  文件不存在: {file_path}")
    
    if not test_files:
        print("❌ 没有找到可用的测试文件")
        return False
    
    # 测试上传
    success_count = 0
    for file_path in test_files:
        print(f"\n📤 测试上传: {file_path}")
        
        try:
            with open(file_path, 'rb') as f:
                files = {'file': (os.path.basename(file_path), f, 'application/octet-stream')}
                
                # 测试询价表上传
                print(f"  🔄 上传到询价表...")
                response = requests.post(
                    f"{base_url}/api/upload/inquiry",
                    files=files,
                    auth=("admin", "admin123"),
                    timeout=30
                )
                
                if response.status_code == 200:
                    result = response.json()
                    print(f"  ✅ 询价表上传成功: {result.get('filename', 'N/A')}")
                    success_count += 1
                else:
                    print(f"  ❌ 询价表上传失败: {response.status_code}")
                    print(f"     错误信息: {response.text}")
                
                # 如果是图片文件，也测试OCR专用接口
                if file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.bmp', '.tiff')):
                    print(f"  🔄 测试OCR处理...")
                    f.seek(0)  # 重置文件指针
                    response = requests.post(
                        f"{base_url}/api/ocr/process-image",
                        files=files,
                        auth=("admin", "admin123"),
                        timeout=30
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        print(f"  ✅ OCR处理成功:")
                        print(f"     原始文本长度: {len(result.get('original_text', ''))}")
                        print(f"     提取项目数: {result.get('statistics', {}).get('extracted_items', 0)}")
                    else:
                        print(f"  ❌ OCR处理失败: {response.status_code}")
                        print(f"     错误信息: {response.text}")
                
        except Exception as e:
            print(f"  ❌ 上传异常: {e}")
    
    print(f"\n📊 上传测试结果: {success_count}/{len(test_files)} 成功")
    return success_count > 0

def test_server_status():
    """测试服务器状态"""
    print("\n🌐 测试服务器状态")
    print("=" * 50)
    
    base_url = "http://localhost:8001"
    
    try:
        response = requests.get(f"{base_url}/", timeout=5)
        if response.status_code == 200:
            print("✅ 服务器运行正常")
            return True
        else:
            print(f"❌ 服务器响应异常: {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ 无法连接到服务器: {e}")
        print("💡 请确保服务器已启动: uvicorn main:app --reload --host 0.0.0.0 --port 8001")
        return False

def main():
    """主测试函数"""
    print("🚀 开始全面文件上传测试")
    print("=" * 60)
    
    # 1. 测试服务器状态
    if not test_server_status():
        return
    
    # 2. 测试OCR环境
    ocr_ok = test_ocr_environment()
    
    # 3. 测试Word处理
    word_ok = test_word_processing()
    
    # 4. 测试文件上传API
    upload_ok = test_file_upload_api()
    
    # 总结
    print("\n" + "=" * 60)
    print("📋 测试总结")
    print("=" * 60)
    print(f"服务器状态: {'✅ 正常' if test_server_status() else '❌ 异常'}")
    print(f"OCR环境: {'✅ 正常' if ocr_ok else '❌ 异常'}")
    print(f"Word处理: {'✅ 正常' if word_ok else '❌ 异常'}")
    print(f"文件上传: {'✅ 正常' if upload_ok else '❌ 异常'}")
    
    if not ocr_ok:
        print("\n🔧 OCR环境问题建议:")
        print("1. 检查 tesseract.exe 路径是否正确")
        print("2. 设置 TESSDATA_PREFIX 环境变量")
        print("3. 确保 tesseract.exe 有执行权限")
    
    if not word_ok:
        print("\n🔧 Word处理问题建议:")
        print("1. 检查 python-docx 包是否正确安装")
        print("2. 确保Word文件包含标准表格")
    
    if not upload_ok:
        print("\n🔧 文件上传问题建议:")
        print("1. 检查服务器日志获取详细错误信息")
        print("2. 确保文件格式正确")
        print("3. 检查网络连接和认证")

if __name__ == "__main__":
    main() 