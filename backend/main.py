from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.openapi.docs import get_swagger_ui_html
from fastapi.openapi.utils import get_openapi
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import sys
import shutil
import secrets
import subprocess
from pathlib import Path
from datetime import datetime
import pandas as pd
import uuid
import time
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
# 导入OCR配置模块
from ocr_config import setup_ocr_environment

# 设置OCR环境
setup_ocr_environment()
import io

# 将当前目录加入Python路径
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# 导入现有的处理脚本
from convert_excel_to_csv import process_excel_to_standard_csv, extract_valve_info
from valve_model_generator import generate_valve_models, analyze_valve_missing_params, parse_valve_info, parse_valve_info_from_combined
from generate_quotes import process_inquiry_file, generate_summary_report
from default_rules import DefaultRulesManager
from csv_utils import safe_read_csv, safe_to_csv
from ocr_correction import OCRCorrector
from enhanced_quote_processor import process_quote_with_enhanced_matching, generate_multi_brand_quote
from structured_quote_generator import generate_structured_quote

# ------------------ OpenAPI / Swagger 配置 ------------------
tags_metadata = [
    {"name": "auth", "description": "账户登录与管理相关接口"},
    {"name": "rules", "description": "默认规则获取与保存"},
    {"name": "interactive", "description": "交互式型号补全与报价流程"},
    {"name": "upload", "description": "价格表 / 询价表 / OCR 图片上传解析"},
    {"name": "ocr", "description": "OCR 识别与文本纠错"},
    {"name": "files", "description": "用户文件列出与生成结果"}
]

app = FastAPI(
    title="阀门报价系统 API",
    description="用于阀门询价解析、型号生成、价格匹配与报价单生成的后端接口。\n\n" \
                "包含：\n" \
                "1. 账户管理与认证 (HTTP Basic)\n" \
                "2. 询价 / 价格表文件上传解析\n" \
                "3. 智能型号补全 + 交互式参数选择\n" \
                "4. 报价生成与汇总\n" \
                "5. OCR 图片识别及纠错\n",
    version="1.0.0",
    openapi_url="/openapi.json",
    docs_url="/docs",          # 默认 Swagger UI
    redoc_url="/redoc",         # ReDoc 文档
    openapi_tags=tags_metadata,
    swagger_ui_parameters={
        "defaultModelsExpandDepth": 0,
        "displayRequestDuration": True,
        "tryItOutEnabled": True
    }
)


# CORS配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8000", "http://127.0.0.1:8000"],  # 指定具体的前端源地址
    allow_credentials=True,
    allow_methods=["*"],  # 允许所有HTTP方法
    allow_headers=["*"],  # 允许所有HTTP头
    expose_headers=["*"],  # 暴露所有头信息
    max_age=86400,  # 预检请求缓存时间（秒）
)

# 基础认证
security = HTTPBasic()

# 账户文件路径
ACCOUNTS_FILE = os.path.join(current_dir, "accounts.txt")

# 数据存储根目录 - 使用quote-system目录下的merchant_data
DATA_ROOT = os.path.abspath(os.path.join(current_dir, "merchant_data"))
print(f"🔧 [INIT] 数据存储根目录: {DATA_ROOT}")

# 确保目录存在
os.makedirs(DATA_ROOT, exist_ok=True)

# 交互式批次数据存储（内存缓存）
interactive_batches = {}

def get_rules_manager():
    """获取规则管理器实例"""
    # 获取当前脚本所在目录的绝对路径
    current_dir = os.path.dirname(os.path.abspath(__file__))
    data_root = os.path.abspath(os.path.join(current_dir, "merchant_data"))
    return DefaultRulesManager(data_root)

class Account(BaseModel):
    username: str
    password: str

class DefaultRules(BaseModel):
    product_defaults: Dict[str, Dict[str, str]]
    custom_products: Dict[str, Dict[str, str]]

class InteractiveChoice(BaseModel):
    valve_info: Dict[str, Any]
    selections: Dict[str, str]

# 新增交互式批量处理相关的数据模型
class InteractiveBatch(BaseModel):
    batch_id: str
    incomplete_items: List[Dict[str, Any]]
    completed_items: List[Dict[str, Any]]
    current_index: int
    total_count: int

class InteractiveSelection(BaseModel):
    batch_id: str
    item_index: int
    selections: Dict[str, str]

class InteractiveItem(BaseModel):
    index: int
    name: str
    specs: str
    quantity: str
    missing_params: List[str]
    valve_info: Dict[str, Any]

# ------------- 通用/响应模型 -------------
class MessageResponse(BaseModel):
    message: str

class LoginResponse(BaseModel):
    username: str
    is_admin: bool

class CreateAccountResponse(MessageResponse):
    pass

class FileListResponse(BaseModel):
    price_tables: List[str]
    inquiry_tables: List[str]
    quotes: List[str]

class OCRProcessResponse(BaseModel):
    message: str
    original_text: str
    corrected_text: str
    extracted_data: List[Dict[str, Any]]
    statistics: Dict[str, Any]

class UploadResponse(BaseModel):
    message: str
    filename: str

class PriceUploadResponse(UploadResponse):
    brands: List[str]

class InteractiveStartProgress(BaseModel):
    current: int
    total: int

class InteractiveStartResponse(BaseModel):
    need_interaction: bool
    batch_id: str
    message: Optional[str] = None
    current_item: Optional[InteractiveItem] = None
    progress: Optional[InteractiveStartProgress] = None

def verify_credentials(credentials: HTTPBasicCredentials = Depends(security)):
    """验证用户凭据"""
    if not os.path.exists(ACCOUNTS_FILE):
        raise HTTPException(status_code=401, detail="账户系统未初始化")
    
    with open(ACCOUNTS_FILE, 'r') as f:
        for line in f:
            line = line.strip()
            if line and ':' in line:
                try:
                    stored_user, stored_pass = line.split(':', 1)
                    if credentials.username == stored_user and credentials.password == stored_pass:
                        return credentials.username
                except ValueError:
                    continue
    
    raise HTTPException(status_code=401, detail="用户名或密码错误")

# 可选：受保护的文档入口（访问 /secure-docs 需登录）
@app.get("/secure-docs", include_in_schema=False)
def secure_docs(username: str = Depends(verify_credentials)):
    """需要登录才能访问的 Swagger UI 页面"""
    return get_swagger_ui_html(openapi_url="/openapi.json", title="Secure Swagger UI")

def is_admin(username: str) -> bool:
    """检查是否为管理员"""
    return username == "admin"

@app.get("/", response_model=MessageResponse, summary="健康检查 / 根信息", tags=["auth"], description="返回 API 基本可用状态。")
async def root():
    return {"message": "报价系统API 正常运行"}

@app.post("/api/login", response_model=LoginResponse, summary="用户登录", tags=["auth"], description="使用 HTTP Basic 认证登录，成功返回用户名及是否为管理员标记。")
async def login(credentials: HTTPBasicCredentials = Depends(security)):
    username = verify_credentials(credentials)
    return {"username": username, "is_admin": is_admin(username)}

@app.get("/api/accounts", response_model=List[str], summary="列出全部账户", tags=["auth"], description="仅管理员可用，返回系统中所有普通与管理员账号用户名列表。")
async def list_accounts(username: str = Depends(verify_credentials)):
    if not is_admin(username):
        raise HTTPException(status_code=403, detail="无权限")
    
    if not os.path.exists(ACCOUNTS_FILE):
        return []
    
    accounts = []
    with open(ACCOUNTS_FILE, 'r') as f:
        for line in f:
            if ':' in line:
                user, _ = line.strip().split(':', 1)
                accounts.append(user)
    return accounts

@app.post("/api/accounts", response_model=CreateAccountResponse, summary="创建账户", tags=["auth"], description="仅管理员：创建一个新账户，并初始化其默认规则与目录结构。")
async def create_account(account: Account, username: str = Depends(verify_credentials)):
    if not is_admin(username):
        raise HTTPException(status_code=403, detail="无权限")
    
    if ':' in account.username:
        raise HTTPException(status_code=400, detail="用户名不能包含冒号")
    if not account.username.strip():
        raise HTTPException(status_code=400, detail="用户名不能为空")
    if not account.password.strip():
        raise HTTPException(status_code=400, detail="密码不能为空")
    
    # 检查账户是否已存在
    if os.path.exists(ACCOUNTS_FILE):
        with open(ACCOUNTS_FILE, 'r') as f:
            for line in f:
                if ':' in line:
                    stored_user, _ = line.strip().split(':', 1)
                    if stored_user == account.username:
                        raise HTTPException(status_code=400, detail="用户已存在")
    
    # 检查文件是否需要添加换行符
    need_newline = False
    if os.path.exists(ACCOUNTS_FILE) and os.path.getsize(ACCOUNTS_FILE) > 0:
        with open(ACCOUNTS_FILE, 'rb') as f:
            f.seek(-1, 2)
            last_char = f.read(1)
            if last_char != b'\n':
                need_newline = True
    
    # 添加新账户
    with open(ACCOUNTS_FILE, 'a') as f:
        if need_newline:
            f.write('\n')
        f.write(f"{account.username}:{account.password}\n")
    
    # 创建用户数据目录
    user_dir = os.path.join(DATA_ROOT, account.username)
    os.makedirs(user_dir, exist_ok=True)
    os.makedirs(os.path.join(user_dir, "价格表"), exist_ok=True)
    os.makedirs(os.path.join(user_dir, "询价表"), exist_ok=True)
    os.makedirs(os.path.join(user_dir, "报价单"), exist_ok=True)
    
    # 为新用户创建默认规则文件
    rules_manager = get_rules_manager()
    rules_manager.create_default_rules_for_new_user(account.username)
    
    return {"message": "账户创建成功"}

@app.delete("/api/accounts/{account_name}", response_model=MessageResponse, summary="删除账户", tags=["auth"], description="仅管理员：删除指定普通账户及其数据。")
async def delete_account(account_name: str, username: str = Depends(verify_credentials)):
    if not is_admin(username):
        raise HTTPException(status_code=403, detail="无权限")
    
    if account_name == "admin":
        raise HTTPException(status_code=400, detail="不能删除管理员账户")
    
    # 从账户文件中删除
    if os.path.exists(ACCOUNTS_FILE):
        lines = []
        with open(ACCOUNTS_FILE, 'r') as f:
            lines = f.readlines()
        
        with open(ACCOUNTS_FILE, 'w') as f:
            for line in lines:
                if not line.startswith(account_name + ':'):
                    f.write(line)
    
    # 删除用户数据目录
    user_dir = os.path.join(DATA_ROOT, account_name)
    if os.path.exists(user_dir):
        shutil.rmtree(user_dir)
    
    return {"message": "账户删除成功"}

# 默认规则管理接口
@app.get("/api/default-rules/options", summary="获取规则可选项", tags=["rules"], description="返回前端构建下拉选择所需的所有默认规则字段与候选值。")
async def get_rule_options(username: str = Depends(verify_credentials)):
    rules_manager = get_rules_manager()
    return rules_manager.get_options_for_frontend()

@app.get("/api/default-rules", summary="获取当前用户默认规则", tags=["rules"], description="读取并返回当前登录用户保存的默认规则配置。")
async def get_user_default_rules(username: str = Depends(verify_credentials)):
    rules_manager = get_rules_manager()
    print(f"🔍 [API] 获取用户默认规则: username={username}")
    return rules_manager.load_user_rules(username)

@app.post("/api/default-rules", response_model=MessageResponse, summary="保存当前用户默认规则", tags=["rules"], description="覆盖保存当前登录用户的默认规则配置。")
async def save_user_default_rules(rules: DefaultRules, username: str = Depends(verify_credentials)):
    rules_manager = get_rules_manager()
    print(f"💾 [API] 保存用户默认规则: username={username}")
    success = rules_manager.save_user_rules(username, rules.dict())
    if success:
        return {"message": "默认规则保存成功"}
    else:
        raise HTTPException(status_code=500, detail="保存默认规则失败")

@app.post("/api/interactive-match")
async def interactive_match(choice: InteractiveChoice, username: str = Depends(verify_credentials)):
    """交互式匹配处理"""
    rules_manager = get_rules_manager()
    print(f"🔧 [API] 交互式匹配: username={username}")
    
    # 应用用户选择
    valve_info = choice.valve_info.copy()
    valve_info.update(choice.selections)
    
    # 使用默认规则管理器处理
    completed_info = rules_manager.apply_default_rules(username, valve_info)
    
    return {
        "message": "匹配完成",
        "valve_info": completed_info
    }

@app.post("/api/get-interactive-options")
async def get_interactive_options(valve_info: Dict[str, Any], username: str = Depends(verify_credentials)):
    """获取交互式选择选项"""
    rules_manager = get_rules_manager()
    options = rules_manager.create_interactive_options(valve_info)
    return options

@app.post("/api/start-interactive-quote", response_model=InteractiveStartResponse, summary="启动交互式报价分析", tags=["interactive"], description="解析询价表并找出需要人工补全参数的产品，返回第一项待交互条目。全部完整则直接标记无需交互。")
async def start_interactive_quote(
    price_file: str = Form(...),
    inquiry_file: str = Form(...),
    username: str = Depends(verify_credentials)
):
    """开始交互式报价流程，分析询价表并返回需要用户选择的产品列表"""
    # 确保在一个有效的目录中运行
    try:
        original_dir = os.getcwd()
    except FileNotFoundError:
        # 如果当前目录不存在，切换到脚本所在目录
        original_dir = os.path.dirname(os.path.abspath(__file__))
        os.chdir(original_dir)
        print(f"🔄 [INTERACTIVE] 切换到有效目录: {original_dir}")
    
    work_dir = None
    
    try:
        print(f"🚀 [INTERACTIVE] 开始交互式报价流程")
        print(f"📁 [INTERACTIVE] 用户: {username}")
        print(f"💰 [INTERACTIVE] 价格文件: {price_file}")
        print(f"📋 [INTERACTIVE] 询价文件: {inquiry_file}")
        
        user_dir = os.path.join(DATA_ROOT, username)
        inquiry_path = os.path.join(user_dir, "询价表", inquiry_file)
        
        # 检查询价文件是否存在
        if not os.path.exists(inquiry_path):
            raise HTTPException(status_code=404, detail=f"询价文件不存在: {inquiry_file}")
        
        # 创建临时工作目录
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        work_dir = os.path.join(user_dir, f"temp_interactive_{timestamp}")
        os.makedirs(work_dir, exist_ok=True)
        
        try:
            # 复制询价文件到工作目录
            inquiry_dst = os.path.join(work_dir, "客户询价表")
            os.makedirs(inquiry_dst, exist_ok=True)
            inquiry_file_dst = os.path.join(inquiry_dst, inquiry_file)
            shutil.copy2(inquiry_path, inquiry_file_dst)
            
            # 切换到工作目录
            os.chdir(work_dir)
            
            # 创建标准格式目录
            os.makedirs("规范后客户询价表数据", exist_ok=True)
            
            # 转换文件格式
            if inquiry_file.endswith(('.xlsx', '.xls')):
                # Excel文件转换
                input_file = os.path.join("客户询价表", inquiry_file)
                output_file = os.path.join("规范后客户询价表数据", f"{Path(inquiry_file).stem}_标准格式.csv")
                process_excel_to_standard_csv(input_file, output_file)
            else:
                # CSV文件直接复制
                input_file = os.path.join("客户询价表", inquiry_file)
                output_file = os.path.join("规范后客户询价表数据", f"{Path(inquiry_file).stem}_标准格式.csv")
                shutil.copy2(input_file, output_file)
            
            # 读取标准格式的询价表
            csv_files = [f for f in os.listdir("规范后客户询价表数据") if f.endswith('.csv')]
            if not csv_files:
                raise ValueError("没有找到转换后的CSV文件")
            
            # 分析第一个CSV文件中的产品
            csv_file = csv_files[0]
            csv_path = os.path.join("规范后客户询价表数据", csv_file)
            df = safe_read_csv(csv_path)
            
            print(f"📊 [INTERACTIVE] 读取询价表: {len(df)} 行数据")
            
            # 分析每个产品，找出需要交互选择的产品
            incomplete_items = []
            completed_items = []
            
            for index, row in df.iterrows():
                if pd.isna(row['品名']) or row['品名'] == '合计':
                    continue
                
                print(f"🔍 [INTERACTIVE] 分析第 {index+1} 行: {row['品名']}")
                
                # 分析缺失参数
                analysis_result = analyze_valve_missing_params(row['品名'], row['规格型号'])
                
                if analysis_result:
                    # 需要交互选择
                    item = {
                        'index': index,
                        'name': str(row['品名']),
                        'specs': str(row['规格型号']),
                        'quantity': str(row.get('数量', '')),
                        'missing_params': analysis_result['missing_params'],
                        'valve_info': analysis_result['valve_info']
                    }
                    incomplete_items.append(item)
                    print(f"❓ [INTERACTIVE] 需要交互: {item['name']} - 缺失参数: {item['missing_params']}")
                else:
                    # 不需要交互选择
                    completed_items.append({
                        'index': index,
                        'name': str(row['品名']),
                        'specs': str(row['规格型号']),
                        'quantity': str(row.get('数量', ''))
                    })
                    print(f"✅ [INTERACTIVE] 无需交互: {row['品名']}")
            
            # 生成批次ID
            batch_id = str(uuid.uuid4())
            
            # 创建批次数据
            batch_data = {
                'batch_id': batch_id,
                'username': username,
                'price_file': price_file,
                'inquiry_file': inquiry_file,
                'work_dir': work_dir,
                'original_dir': original_dir,
                'csv_file': csv_file,
                'incomplete_items': incomplete_items,
                'completed_items': completed_items,
                'current_index': 0,
                'total_count': len(incomplete_items),
                'user_selections': {}  # 存储用户的选择
            }
            
            # 存储批次数据
            interactive_batches[batch_id] = batch_data
            
            print(f"📊 [INTERACTIVE] 分析完成:")
            print(f"   需要交互的产品: {len(incomplete_items)} 个")
            print(f"   无需交互的产品: {len(completed_items)} 个")
            print(f"   批次ID: {batch_id}")
            
            if not incomplete_items:
                # 没有需要交互的产品，直接生成报价
                return {
                    "need_interaction": False,
                    "message": "所有产品参数完整，无需交互选择",
                    "batch_id": batch_id
                }
            
            # 返回第一个需要交互的产品
            first_item = incomplete_items[0]
            return {
                "need_interaction": True,
                "batch_id": batch_id,
                "current_item": first_item,
                "progress": {
                    "current": 1,
                    "total": len(incomplete_items)
                }
            }
            
        except Exception as e:
            # 清理工作目录
            os.chdir(original_dir)
            if work_dir and os.path.exists(work_dir):
                shutil.rmtree(work_dir)
            raise e
            
    except Exception as e:
        # 确保返回到原始目录
        try:
            os.chdir(original_dir)
        except Exception as cd_error:
            print(f"⚠️  [INTERACTIVE] 无法返回原始目录: {cd_error}")
            # 尝试切换到脚本所在目录
            script_dir = os.path.dirname(os.path.abspath(__file__))
            try:
                os.chdir(script_dir)
                print(f"🔄 [INTERACTIVE] 切换到脚本目录: {script_dir}")
            except Exception:
                pass
        
        # 清理可能创建的工作目录
        if work_dir and os.path.exists(work_dir):
            try:
                shutil.rmtree(work_dir)
                print(f"🗑️  [INTERACTIVE] 清理临时目录: {work_dir}")
            except Exception as cleanup_error:
                print(f"⚠️  [INTERACTIVE] 清理临时目录失败: {cleanup_error}")
                
        print(f"❌ [INTERACTIVE] 启动交互式流程失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"启动交互式流程失败: {str(e)}")

@app.post("/api/submit-interactive-selection")
async def submit_interactive_selection(
    selection: InteractiveSelection,
    username: str = Depends(verify_credentials)
):
    """提交单个产品的参数选择"""
    try:
        print(f"📝 [INTERACTIVE] 提交参数选择: batch_id={selection.batch_id}")
        
        # 获取批次数据
        if selection.batch_id not in interactive_batches:
            raise HTTPException(status_code=404, detail="批次不存在或已过期")
        
        batch_data = interactive_batches[selection.batch_id]
        
        # 验证用户权限
        if batch_data['username'] != username:
            raise HTTPException(status_code=403, detail="无权限访问此批次")
        
        # 验证项目索引
        if selection.item_index >= len(batch_data['incomplete_items']):
            raise HTTPException(status_code=400, detail="无效的项目索引")
        
        # 保存用户选择
        item_key = f"item_{selection.item_index}"
        batch_data['user_selections'][item_key] = selection.selections
        
        print(f"✅ [INTERACTIVE] 保存选择: {selection.selections}")
        
        # 更新当前索引
        batch_data['current_index'] = selection.item_index + 1
        
        # 检查是否还有未完成的项目
        if batch_data['current_index'] < len(batch_data['incomplete_items']):
            # 返回下一个需要交互的产品
            next_item = batch_data['incomplete_items'][batch_data['current_index']]
            return {
                "completed": False,
                "next_item": next_item,
                "progress": {
                    "current": batch_data['current_index'] + 1,
                    "total": len(batch_data['incomplete_items'])
                }
            }
        else:
            # 所有交互完成
            return {
                "completed": True,
                "message": "所有参数选择完成，可以生成报价单",
                "total_selections": len(batch_data['user_selections'])
            }
            
    except Exception as e:
        print(f"❌ [INTERACTIVE] 提交选择失败: {e}")
        raise HTTPException(status_code=500, detail=f"提交选择失败: {str(e)}")

@app.post("/api/complete-interactive-quote")
async def complete_interactive_quote(
    batch_id: str = Form(...),
    username: str = Depends(verify_credentials)
):
    """完成所有交互选择后，生成最终报价单"""
    try:
        print(f"🎯 [INTERACTIVE] 完成交互式报价: batch_id={batch_id}")
        
        # 获取批次数据
        if batch_id not in interactive_batches:
            raise HTTPException(status_code=404, detail="批次不存在或已过期")
        
        batch_data = interactive_batches[batch_id]
        
        # 验证用户权限
        if batch_data['username'] != username:
            raise HTTPException(status_code=403, detail="无权限访问此批次")
        
        # 获取原始目录和工作目录
        original_dir = batch_data.get('original_dir')
        work_dir = batch_data.get('work_dir')
        
        # 验证目录是否存在
        if not original_dir or not os.path.exists(original_dir):
            # 如果原始目录不存在，使用脚本所在目录
            original_dir = os.path.dirname(os.path.abspath(__file__))
            print(f"🔄 [INTERACTIVE] 原始目录不存在，使用脚本目录: {original_dir}")
        
        if not work_dir or not os.path.exists(work_dir):
            raise HTTPException(status_code=500, detail="工作目录不存在，无法完成报价")
        
        # 切换到工作目录
        current_dir = os.getcwd()
        os.chdir(work_dir)
        print(f"🔄 [INTERACTIVE] 从 {current_dir} 切换到工作目录: {work_dir}")
        
        try:
            from valve_model_generator import parse_valve_info
            
            # 创建型号编码目录
            os.makedirs("型号编码后的询价表数据", exist_ok=True)
            
            # 读取原始询价表
            csv_path = os.path.join("规范后客户询价表数据", batch_data['csv_file'])
            df = safe_read_csv(csv_path)
            
            # 生成型号列
            models = []
            for index, row in df.iterrows():
                if pd.isna(row['品名']) or row['品名'] == '合计':
                    models.append('')
                    continue
                
                # 检查是否是需要交互的产品
                item_key = f"item_{index}"
                if item_key in batch_data['user_selections']:
                    # 使用用户选择的参数
                    print(f"🔧 [INTERACTIVE] 使用用户选择生成型号: 第{index+1}行")
                    
                    # 获取基础阀门信息
                    for item in batch_data['incomplete_items']:
                        if item['index'] == index:
                            valve_info = item['valve_info'].copy()
                            user_selections = batch_data['user_selections'][item_key]
                            
                            # 应用用户选择
                            valve_info.update(user_selections)
                            
                            # 生成型号
                            model = generate_model_from_valve_info(valve_info)
                            models.append(model)
                            print(f"✅ [INTERACTIVE] 生成型号: {model}")
                            break
                    else:
                        # 找不到对应的项目，使用默认方式
                        model = parse_valve_info(row['品名'], row['规格型号'], username, True)
                        models.append(model)
                else:
                    # 使用默认方式生成型号
                    model = parse_valve_info(row['品名'], row['规格型号'], username, True)
                    models.append(model)
            
            # 添加型号列
            df['标准型号'] = models
            
            # 保存型号编码后的文件
            output_file = os.path.join("型号编码后的询价表数据", batch_data['csv_file'])
            safe_to_csv(df, output_file)
            
            print(f"✅ [INTERACTIVE] 型号生成完成，保存到: {output_file}")
            
            # 继续后续的报价生成流程
            # 处理价格对照表
            price_file = batch_data['price_file']
            # 使用绝对路径
            price_path = os.path.join(DATA_ROOT, username, "价格表", price_file)
            
            os.makedirs("规范后的价格对照表数据", exist_ok=True)
            price_csv_path = os.path.join("规范后的价格对照表数据", "价格.csv")
            
            if price_file.endswith('.csv'):
                price_df = safe_read_csv(price_path)
            else:
                price_df = pd.read_excel(price_path)
            
            # 检查和映射价格表列名
            required_columns = ['型号', '规格', '品牌', '价格']
            column_mapping = {}
            for col in price_df.columns:
                col_lower = str(col).lower()
                if '型号' in col_lower:
                    column_mapping[col] = '型号'
                elif '规格' in col_lower or 'dn' in col_lower:
                    column_mapping[col] = '规格'
                elif '品牌' in col_lower:
                    column_mapping[col] = '品牌'
                elif '价格' in col_lower or '单价' in col_lower:
                    column_mapping[col] = '价格'
            
            if column_mapping:
                price_df = price_df.rename(columns=column_mapping)
            
            price_df.to_csv(price_csv_path, index=False, encoding='utf-8-sig')
            
            # 生成报价
            from generate_quotes import process_inquiry_file, generate_summary_report
            
            os.makedirs("报价数据", exist_ok=True)
            
            processed_files = []
            output_file = process_inquiry_file(output_file, price_df)
            if output_file:
                processed_files.append(output_file)
            
            # 生成汇总报告
            if processed_files:
                generate_summary_report(processed_files)
            
            # 复制生成的报价到用户报价目录
            quote_dir = os.path.join(DATA_ROOT, username, "报价单")
            os.makedirs(quote_dir, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            generated_files = []
            
            if os.path.exists("报价数据"):
                for file in os.listdir("报价数据"):
                    if file.endswith(('.xlsx', '.csv')):
                        src = os.path.join("报价数据", file)
                        dst = os.path.join(quote_dir, f"{timestamp}_交互式_{file}")
                        shutil.copy2(src, dst)
                        generated_files.append(f"{timestamp}_交互式_{file}")
            
            print(f"🎉 [INTERACTIVE] 交互式报价完成，生成文件: {generated_files}")
            
            return {
                "message": "交互式报价单生成成功",
                "files": generated_files,
                "total_interactions": len(batch_data['user_selections'])
            }
            
        finally:
            # 返回原目录并清理
            try:
                os.chdir(original_dir)
                print(f"🔄 [INTERACTIVE] 返回原始目录: {original_dir}")
            except Exception as e:
                # 如果无法返回原始目录，尝试切换到脚本目录
                script_dir = os.path.dirname(os.path.abspath(__file__))
                try:
                    os.chdir(script_dir)
                    print(f"🔄 [INTERACTIVE] 无法返回原始目录，切换到脚本目录: {script_dir}")
                except Exception:
                    print(f"⚠️  [INTERACTIVE] 无法切换到任何有效目录")
            
            try:
                if os.path.exists(work_dir):
                    shutil.rmtree(work_dir)
                    print(f"🗑️  [INTERACTIVE] 清理临时目录: {work_dir}")
            except Exception as e:
                print(f"⚠️  [INTERACTIVE] 清理临时目录失败: {e}")
            
            # 清理批次数据
            if batch_id in interactive_batches:
                del interactive_batches[batch_id]
                print(f"🗑️  [INTERACTIVE] 清理批次数据: {batch_id}")
            
    except Exception as e:
        print(f"❌ [INTERACTIVE] 完成交互式报价失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"完成交互式报价失败: {str(e)}")

def generate_model_from_valve_info(valve_info):
    """根据阀门信息生成型号"""
    valve_type = valve_info.get('product_type', '')
    drive_mode = valve_info.get('drive_mode', '')
    connection = valve_info.get('connection', '')
    structure = valve_info.get('structure', '')
    sealing = valve_info.get('sealing', '')
    pressure = valve_info.get('pressure', '16')
    material = valve_info.get('material', 'Q')
    
    # 对于特殊产品，直接返回完整型号
    if valve_type in ['100X', '200X', '500X', '800X']:
        model = ""
        if connection == '8':
            model = "8"
        model += valve_type + f"-{pressure}{material}"
        return model
    
    # 组合标准型号
    model = valve_type
    
    # 驱动方式（手动默认不标）
    if drive_mode:
        model += drive_mode
    
    # 连接方式
    model += connection
    
    # 结构形式
    model += structure
    
    # 密封材料
    model += sealing
    
    # 压力-材料
    model += f"-{pressure}{material}"
    
    return model

def read_csv_with_encoding_fallback(file_bytes):
    import pandas as pd
    import io
    try:
        return pd.read_csv(io.BytesIO(file_bytes))
    except UnicodeDecodeError:
        try:
            return pd.read_csv(io.BytesIO(file_bytes), encoding='gbk')
        except Exception as e:
            raise e

# 文件解析为 Excel 的工具函数
def parse_file_to_excel(file_bytes, filename, save_dir):
    import pandas as pd
    import os
    ext = os.path.splitext(filename)[-1].lower()
    excel_name = os.path.splitext(filename)[0] + ".xlsx"
    save_path = os.path.join(save_dir, excel_name)
    df = None

    print(f"🔍 [PARSE] 开始解析文件: {filename}, 扩展名: {ext}")

    try:
        if ext in [".csv"]:
            print(f"📊 [PARSE] 解析CSV文件")
            df = read_csv_with_encoding_fallback(file_bytes)
        elif ext in [".xlsx", ".xls"]:
            print(f"📊 [PARSE] 解析Excel文件")
            df = pd.read_excel(io.BytesIO(file_bytes))
        elif ext == ".pdf":
            print(f"📊 [PARSE] 解析PDF文件")
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    print(f"📄 [PARSE] 处理PDF第{page_num+1}页")
                    table = page.extract_table()
                    if table:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        print(f"✅ [PARSE] PDF第{page_num+1}页找到表格，共{len(df)}行")
                        break
            if df is None:
                raise Exception("PDF未检测到表格")
        elif ext == ".docx":
            print(f"📊 [PARSE] 解析Word文件")
            doc = Document(io.BytesIO(file_bytes))
            print(f"📄 [PARSE] Word文档包含{len(doc.tables)}个表格")
            for table_num, table in enumerate(doc.tables):
                print(f"📋 [PARSE] 处理Word第{table_num+1}个表格")
                data = [[cell.text for cell in row.cells] for row in table.rows]
                if data and len(data) > 1:  # 确保有表头和数据
                    df = pd.DataFrame(data[1:], columns=data[0])
                    print(f"✅ [PARSE] Word第{table_num+1}个表格解析成功，共{len(df)}行")
                    break
            if df is None:
                raise Exception("Word未检测到有效表格")
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
            print(f"📊 [PARSE] 解析图片文件")
            image = Image.open(io.BytesIO(file_bytes))
            print(f"🖼️ [PARSE] 图片尺寸: {image.size}")
            
            # 使用OCR提取文本
            try:
                text = pytesseract.image_to_string(
                    image,
                    lang='eng'
                )
                print(f"📝 [PARSE] OCR提取文本长度: {len(text)}")
                print(f"📝 [PARSE] OCR提取文本预览: {text[:200]}...")
                
                if not text.strip():
                    raise Exception("图片未检测到有效文本内容")
                
                # 使用OCR修正器处理文本
                corrector = OCRCorrector()
                results = corrector.process_ocr_text(text)
                
                # 创建数据框
                excel_data = []
                for item in results['extracted_data']:
                    excel_data.append({
                        '品名': f"阀门 DN{item['dn_value']}",
                        '规格型号': f"DN{item['dn_value']}",
                        '数量': item['quantity'],
                        '单位': '个',
                        '原始文本': item['original_text'],
                        '修正文本': item['corrected_text'],
                        '置信度': item['confidence']
                    })
                
                # 如果没有提取到数据，创建一个包含原始文本的行
                if not excel_data:
                    excel_data.append({
                        '品名': 'OCR提取文本',
                        '规格型号': '原始文本',
                        '数量': '1',
                        '单位': '个',
                        '原始文本': text,
                        '修正文本': results['corrected_text'],
                        '置信度': 'low'
                    })
                
                df = pd.DataFrame(excel_data)
                print(f"✅ [PARSE] 图片OCR处理完成，生成{len(df)}行数据")
                
            except Exception as ocr_error:
                print(f"❌ [PARSE] OCR处理失败: {ocr_error}")
                # 如果OCR失败，尝试简单的文本分割
                lines = [line.split() for line in text.splitlines() if line.strip()]
                if not lines:
                    raise Exception("图片未检测到有效文本内容")
                if len(set(len(row) for row in lines)) != 1:
                    print(f"⚠️ [PARSE] 图片表格结构不规整，使用简单分割")
                    # 创建简单的数据框
                    df = pd.DataFrame({
                        '原始文本': [text],
                        '处理状态': ['OCR处理失败，请手动检查']
                    })
                else:
                    df = pd.DataFrame(lines)
        else:
            raise Exception(f"不支持的文件类型: {ext}")

        if df is None:
            raise Exception("文件解析失败，未生成有效数据")

        print(f"💾 [PARSE] 保存Excel文件: {save_path}")
        df.to_excel(save_path, index=False)
        print(f"✅ [PARSE] 文件解析完成: {excel_name}")
        return excel_name

    except Exception as e:
        print(f"❌ [PARSE] 文件解析失败: {e}")
        import traceback
        traceback.print_exc()
        raise e

# 修改上传接口，自动转为 Excel
@app.post("/api/upload/price", response_model=PriceUploadResponse, summary="上传价格表", tags=["upload"], description="上传一个 Excel 价格表，验证列格式并保存（会替换旧价格表）。")
async def upload_price_table(file: UploadFile = File(...), username: str = Depends(verify_credentials)):
    try:
        print(f"📤 [UPLOAD] 开始上传价格表: {file.filename}")
        
        # 检查文件
        if not file.filename:
            raise HTTPException(status_code=400, detail="文件名不能为空")
        
        # 检查文件扩展名
        if not file.filename.endswith(('.xlsx', '.xls')):
            raise HTTPException(status_code=400, detail="只支持Excel文件格式(.xlsx, .xls)")
        
        # 检查文件大小
        file_bytes = await file.read()
        if len(file_bytes) == 0:
            raise HTTPException(status_code=400, detail="文件内容为空")
        
        print(f"📊 [UPLOAD] 文件大小: {len(file_bytes)} 字节")
        
        # 临时保存文件进行验证
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name
        
        try:
            # 读取Excel文件进行验证
            import pandas as pd
            df = pd.read_excel(temp_path)
            
            # 导入验证模块
            from price_validator import validate_price_table_format
            
            # 验证价格表格式
            validation_result = validate_price_table_format(df)
            
            if not validation_result['is_valid']:
                # 验证失败，返回错误信息
                error_message = "价格表格式验证失败:\n" + "\n".join(validation_result['errors'])
                raise HTTPException(status_code=400, detail=error_message)
            
            # 验证成功，检查用户是否已有价格表
            user_dir = os.path.join(DATA_ROOT, username, "价格表")
            os.makedirs(user_dir, exist_ok=True)
            print(f"📁 [UPLOAD] 用户目录: {user_dir}")
            
            # 检查是否已有价格表文件
            existing_files = [f for f in os.listdir(user_dir) if f.endswith(('.xlsx', '.xls', '.csv'))]
            if existing_files:
                # 删除现有价格表文件
                for existing_file in existing_files:
                    existing_path = os.path.join(user_dir, existing_file)
                    try:
                        os.remove(existing_path)
                        print(f"🗑️ [UPLOAD] 删除现有价格表: {existing_file}")
                    except Exception as e:
                        print(f"⚠️ [UPLOAD] 删除现有文件失败: {e}")
            
            # 解析文件
            try:
                excel_name = parse_file_to_excel(file_bytes, file.filename, user_dir)
                print(f"✅ [UPLOAD] 价格表上传成功: {excel_name}")
                
                # 返回成功信息和品牌列表
                return {
                    "message": validation_result['message'] + " (已替换原有价格表)",
                    "brands": validation_result['brands'],
                    "filename": excel_name
                }
            except Exception as parse_error:
                print(f"❌ [UPLOAD] 文件解析失败: {parse_error}")
                import traceback
                traceback.print_exc()
                raise HTTPException(status_code=400, detail=f"文件解析失败: {str(parse_error)}")
            
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [UPLOAD] 价格表上传失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"上传失败: {str(e)}")

@app.post("/api/upload/inquiry", response_model=UploadResponse, summary="上传询价表", tags=["upload"], description="上传任意支持格式的询价表文件，解析并规范为 Excel。")
async def upload_inquiry_table(file: UploadFile = File(...), username: str = Depends(verify_credentials)):
    try:
        print(f"📤 [UPLOAD] 开始上传询价表: {file.filename}")
        
        # 检查文件
        if not file.filename:
            raise HTTPException(status_code=400, detail="文件名不能为空")
        
        # 检查文件大小
        file_bytes = await file.read()
        if len(file_bytes) == 0:
            raise HTTPException(status_code=400, detail="文件内容为空")
        
        print(f"📊 [UPLOAD] 文件大小: {len(file_bytes)} 字节")
        
        # 创建用户目录
        user_dir = os.path.join(DATA_ROOT, username, "询价表")
        os.makedirs(user_dir, exist_ok=True)
        print(f"📁 [UPLOAD] 用户目录: {user_dir}")
        
        # 解析文件
        try:
            excel_name = parse_file_to_excel(file_bytes, file.filename, user_dir)
            print(f"✅ [UPLOAD] 询价表上传成功: {excel_name}")
            return {"message": "询价表上传并解析成功", "filename": excel_name}
        except Exception as parse_error:
            print(f"❌ [UPLOAD] 文件解析失败: {parse_error}")
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=400, detail=f"文件解析失败: {str(parse_error)}")
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [UPLOAD] 询价表上传失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"上传失败: {str(e)}")

@app.post("/api/ocr/process-image", response_model=OCRProcessResponse, summary="图片 OCR 识别与纠错", tags=["ocr"], description="上传图片执行 OCR 识别并对文本进行结构化解析与纠错。")
async def process_image_ocr(file: UploadFile = File(...), username: str = Depends(verify_credentials)):
    try:
        print(f"🔍 [OCR] 开始处理图片OCR: {file.filename}")
        
        # 检查文件类型
        if not file.filename or not file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
            raise HTTPException(status_code=400, detail="只支持图片文件格式")
        
        # 读取图片文件
        file_bytes = await file.read()
        
        # 使用OCR提取文本
        try:
            image = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(
                image,
                lang='eng'
            )
            print(f"📝 [OCR] 原始提取文本: {text[:200]}...")
        except Exception as ocr_error:
            print(f"❌ [OCR] OCR提取失败: {ocr_error}")
            raise HTTPException(status_code=500, detail=f"OCR文本提取失败: {str(ocr_error)}")
        
        # 使用OCR修正器处理文本
        corrector = OCRCorrector()
        results = corrector.process_ocr_text(text)
        
        print(f"✅ [OCR] OCR处理完成:")
        print(f"   原始文本行数: {results['statistics']['original_lines']}")
        print(f"   修正字符数: {results['statistics']['corrections_made']}")
        print(f"   提取项目数: {results['statistics']['extracted_items']}")
        
        return {
            "message": "图片OCR处理成功",
            "original_text": text,
            "corrected_text": results['corrected_text'],
            "extracted_data": results['extracted_data'],
            "statistics": results['statistics']
        }
        
    except Exception as e:
        print(f"❌ [OCR] 图片OCR处理失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"图片OCR处理失败: {str(e)}")

@app.post("/api/ocr/process-image-to-excel", summary="图片 OCR 生成询价 Excel", tags=["ocr"], description="对图片执行 OCR 并直接生成标准结构的询价 Excel 文件存入用户目录。")
async def process_image_to_excel(file: UploadFile = File(...), username: str = Depends(verify_credentials)):
    try:
        print(f"🔍 [OCR] 开始处理图片OCR并生成Excel: {file.filename}")
        
        # 检查文件类型
        if not file.filename or not file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
            raise HTTPException(status_code=400, detail="只支持图片文件格式")
        
        # 读取图片文件
        file_bytes = await file.read()
        
        # 使用OCR提取文本
        try:
            image = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(
                image,
                lang='eng'
            )
            print(f"📝 [OCR] 原始提取文本: {text[:200]}...")
        except Exception as ocr_error:
            print(f"❌ [OCR] OCR提取失败: {ocr_error}")
            raise HTTPException(status_code=500, detail=f"OCR文本提取失败: {str(ocr_error)}")
        
        # 使用OCR修正器处理文本
        corrector = OCRCorrector()
        results = corrector.process_ocr_text(text)
        
        # 生成Excel文件
        import pandas as pd
        from datetime import datetime
        
        # 创建数据框
        excel_data = []
        for item in results['extracted_data']:
            excel_data.append({
                '品名': f"阀门 DN{item['dn_value']}",
                '规格型号': f"DN{item['dn_value']}",
                '数量': item['quantity'],
                '单位': '个',
                '原始文本': item['original_text'],
                '修正文本': item['corrected_text'],
                '置信度': item['confidence']
            })
        
        # 如果没有提取到数据，创建一个包含原始文本的行
        if not excel_data:
            excel_data.append({
                '品名': 'OCR提取文本',
                '规格型号': '原始文本',
                '数量': '1',
                '单位': '个',
                '原始文本': text,
                '修正文本': results['corrected_text'],
                '置信度': 'low'
            })
        
        df = pd.DataFrame(excel_data)
        
        # 保存到用户目录
        user_dir = os.path.join(DATA_ROOT, username, "询价表")
        os.makedirs(user_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_filename = f"OCR_{timestamp}_{Path(file.filename).stem}.xlsx"
        excel_path = os.path.join(user_dir, excel_filename)
        
        df.to_excel(excel_path, index=False)
        
        print(f"✅ [OCR] Excel文件生成成功: {excel_filename}")
        
        return {
            "message": "图片OCR处理并生成Excel成功",
            "filename": excel_filename,
            "original_text": text,
            "corrected_text": results['corrected_text'],
            "extracted_data": results['extracted_data'],
            "statistics": results['statistics']
        }
        
    except Exception as e:
        print(f"❌ [OCR] 图片OCR处理失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"图片OCR处理失败: {str(e)}")

@app.get("/api/files", response_model=FileListResponse, summary="列出用户文件", tags=["files"], description="列出当前用户上传的价格表、询价表以及生成的报价单文件名。")
async def list_files(username: str = Depends(verify_credentials)):
    print(f"📂 获取文件列表请求: username={username}")
    
    user_dir = os.path.join(DATA_ROOT, username)
    files = {
        "price_tables": [],
        "inquiry_tables": [],
        "quotes": []
    }
    
    # 价格表
    price_dir = os.path.join(user_dir, "价格表")
    if os.path.exists(price_dir):
        price_files = []
        for f in os.listdir(price_dir):
            if f.endswith(('.xlsx', '.xls', '.csv')):
                file_path = os.path.join(price_dir, f)
                if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                    price_files.append(f)
        files["price_tables"] = sorted(price_files, key=lambda x: os.path.getmtime(os.path.join(price_dir, x)), reverse=True)
        print(f"📋 价格表文件 ({len(price_files)}个): {price_files}")
    
    # 询价表
    inquiry_dir = os.path.join(user_dir, "询价表")
    if os.path.exists(inquiry_dir):
        inquiry_files = []
        for f in os.listdir(inquiry_dir):
            if f.endswith(('.xlsx', '.xls', '.csv')):
                file_path = os.path.join(inquiry_dir, f)
                if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                    inquiry_files.append(f)
        files["inquiry_tables"] = sorted(inquiry_files, key=lambda x: os.path.getmtime(os.path.join(inquiry_dir, x)), reverse=True)
        print(f"📋 询价表文件 ({len(inquiry_files)}个): {inquiry_files}")
    
    # 报价单
    quote_dir = os.path.join(user_dir, "报价单")
    if os.path.exists(quote_dir):
        quote_files = []
        for f in os.listdir(quote_dir):
            if f.endswith(('.xlsx', '.csv')):
                file_path = os.path.join(quote_dir, f)
                if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                    quote_files.append(f)
        files["quotes"] = sorted(quote_files, key=lambda x: os.path.getmtime(os.path.join(quote_dir, x)), reverse=True)
        print(f"📋 报价单文件 ({len(quote_files)}个): {quote_files}")
    else:
        print(f"❌ 报价单目录不存在: {quote_dir}")
    
    print(f"📊 返回文件列表: {files}")
    return files

def append_quote_to_original(price_file, inquiry_file, output_file, price_columns=None):
    """
    只在原始表格后追加"标准型号、价格、品牌、总价"4列，原内容和格式全部保留。
    自动识别价格表的产品名称、型号、规格字段，提升兼容性。
    """
    orig_df = pd.read_excel(inquiry_file)
    price_df = pd.read_excel(price_file)
    if price_columns:
        price_df = price_df.rename(columns=price_columns)
    orig_df['标准型号'] = ''
    orig_df['价格'] = ''
    orig_df['品牌'] = ''
    orig_df['总价'] = ''

    # 自动识别价格表的产品名称列
    product_name_col = None
    for col in ['产品名称', '名称', '品名', '项目名称', '物料名称']:
        if col in price_df.columns:
            product_name_col = col
            break
    if not product_name_col:
        raise Exception("价格表缺少产品名称相关列")

    # 自动识别价格表的型号列
    model_col = None
    for col in ['型号', '产品型号', '规格型号', '型号编码']:
        if col in price_df.columns:
            model_col = col
            break
    # 型号列可选

    # 自动识别价格表的规格列
    spec_col = None
    for col in ['规格', '规格型号', '口径', '公称直径']:
        if col in price_df.columns:
            spec_col = col
            break
    # 规格列可选

    # 自动识别品牌列
    brand_col = None
    for col in ['品牌', '厂商', '生产厂家']:
        if col in price_df.columns:
            brand_col = col
            break

    for idx, row in orig_df.iterrows():
        name = row.get('名称', '')
        spec = row.get('规格') or row.get('规格型号') or ''
        # 生成标准型号
        orig_df.at[idx, '标准型号'] = parse_valve_info(name, spec)
        # 匹配价格表
        if spec_col:
            match = price_df[
                (price_df[product_name_col] == name) &
                (price_df[spec_col] == spec)
            ]
        else:
            match = price_df[
                (price_df[product_name_col] == name)
            ]
        if not match.empty:
            orig_df.at[idx, '价格'] = match.iloc[0].get('单价', '')
            if brand_col:
                orig_df.at[idx, '品牌'] = match.iloc[0].get(brand_col, '')
            try:
                qty = float(row.get('数量', 1))
                orig_df.at[idx, '总价'] = float(match.iloc[0].get('单价', 0)) * qty
            except Exception:
                orig_df.at[idx, '总价'] = ''
        else:
            orig_df.at[idx, '价格'] = ''
            orig_df.at[idx, '品牌'] = ''
            orig_df.at[idx, '总价'] = ''
    orig_df.to_csv(output_file, index=False, encoding='utf-8-sig')
    return output_file

@app.get("/api/get-brands")
async def get_brands(price_file: str = Query(...), username: str = Depends(verify_credentials)):
    """获取价格表中的品牌列表"""
    try:
        print(f"🔍 [BRANDS] 开始获取品牌列表")
        print(f"📁 [BRANDS] 用户: {username}")
        print(f"💰 [BRANDS] 价格文件: {price_file}")
        
        user_dir = os.path.join(DATA_ROOT, username)
        price_path = os.path.join(user_dir, "价格表", price_file)
        
        if not os.path.exists(price_path):
            raise HTTPException(status_code=404, detail=f"价格文件不存在: {price_file}")
        
        # 导入品牌提取函数
        from price_validator import extract_brands_from_price_table
        
        # 提取品牌列表
        brands = extract_brands_from_price_table(price_path)
        
        print(f"✅ [BRANDS] 成功提取品牌: {brands}")
        return {"brands": brands}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [BRANDS] 获取品牌列表失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"获取品牌列表失败: {str(e)}")

@app.post("/api/generate-quote")
async def generate_quote(
    price_file: str = Form(...),
    inquiry_file: str = Form(...),
    company: str = Form(...),  # 接收前端推断的公司名
    brand: str = Form(None),   # 新增，前端传递的品牌
    scheme: str = Form("scheme1"),  # 方案选择：scheme1=第一种，scheme2=第二种
    username: str = Depends(verify_credentials),
    auto_fill_price: bool = Form(False),  # 新增参数，是否自动用价格表再次填充价格和品牌
    # 第二方案的公司信息参数
    company_name: str = Form(None),
    business_contact: str = Form(None),
    contact_phone: str = Form(None),
    contact_email: str = Form(None),
    customer_header: str = Form(None),
    recipient: str = Form(None),
    contact_method: str = Form(None),
    address: str = Form(None),
    tax_rate: str = Form(None)
):
    try:
        print(f"🚀 [QUOTE] 开始生成价格后的报价单")
        print(f"📁 [QUOTE] 用户: {username}")
        print(f"💰 [QUOTE] 价格文件: {price_file}")
        print(f"📋 [QUOTE] 询价文件: {inquiry_file}")
        print(f"🏢 [QUOTE] 公司: {company}")
        print(f"🏷️ [QUOTE] 品牌: {brand}")
        
        # 检查文件类型，如果是图片或文本文件，强制使用第二种方案
        file_ext = os.path.splitext(inquiry_file)[-1].lower()
        image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif']
        text_extensions = ['.txt', '.doc', '.docx', '.pdf']
        
        if file_ext in image_extensions or file_ext in text_extensions:
            print(f"📸 [QUOTE] 检测到图片/文本文件 ({file_ext})，强制使用第二种方案")
            scheme = "scheme2"
        
        user_dir = os.path.join(DATA_ROOT, username)
        price_path = os.path.join(user_dir, "价格表", price_file)
        inquiry_path = os.path.join(user_dir, "询价表", inquiry_file)
        if not os.path.exists(price_path):
            raise HTTPException(status_code=404, detail=f"价格文件不存在: {price_file}")
        if not os.path.exists(inquiry_path):
            raise HTTPException(status_code=404, detail=f"询价文件不存在: {inquiry_file}")
        base_name = os.path.splitext(os.path.basename(inquiry_file))[0]
        quote_dir = os.path.join(user_dir, "报价单")
        os.makedirs(quote_dir, exist_ok=True)
        standard_csv = os.path.join(quote_dir, f"{base_name}_标准格式.csv")
        standard_xlsx = os.path.join(quote_dir, f"{base_name}_标准格式.xlsx")
        from convert_excel_to_csv import process_excel_to_standard_csv
        # 让标准化流程返回实际生成的文件名（带时间戳）
        result = process_excel_to_standard_csv(inquiry_path, None, price_file=price_path, selected_brand=brand)
        if isinstance(result, tuple):
            standard_csv, standard_xlsx = result
            print(f"[DEBUG] 实际生成的CSV文件: {standard_csv}")
            print(f"[DEBUG] 实际生成的XLSX文件: {standard_xlsx}")
        else:
            # 兼容旧版本返回值
            standard_csv = result
            standard_xlsx = standard_csv.replace('.csv', '.xlsx')
            print(f"[DEBUG] 实际生成的CSV文件: {standard_csv}")
            import pandas as pd
            df = pd.read_csv(standard_csv)
            df.to_excel(standard_xlsx, index=False)
        # 读取数据用于后续处理
        import pandas as pd
        df = pd.read_csv(standard_csv)
        print(f"📊 [QUOTE] 询价表数据读取完成，共{len(df)}行")
        # 读取用户折扣
        rules_manager = get_rules_manager()
        user_discount = rules_manager.get_user_discount(username)
        print(f"[QUOTE] 应用用户折扣: {user_discount}")

        # 根据方案生成
        result_payload = {"message": "生成成功"}
        if scheme == "scheme1":
            # 自动用价格表再次填充价格和品牌（可选）
            if auto_fill_price:
                print(f"[QUOTE] 自动用价格表再次填充价格和品牌...")
                match_quote_with_price_table(standard_xlsx, price_path, brand)
            # 对 standard_xlsx 应用折扣（单价×折扣，总价随之变更）
            try:
                import pandas as pd
                df1 = pd.read_excel(standard_xlsx)
                if '单价' in df1.columns:
                    df1['单价'] = pd.to_numeric(df1['单价'], errors='coerce') * user_discount
                    if '数量' in df1.columns:
                        qty = pd.to_numeric(df1['数量'], errors='coerce').fillna(0)
                        df1['总价'] = (df1['单价'].fillna(0) * qty).round(2)
                    df1.to_excel(standard_xlsx, index=False)
            except Exception as _:
                pass
            print(f"🎉 [QUOTE] 报价单生成成功: {os.path.basename(standard_xlsx)}")
            result_payload.update({"file": os.path.basename(standard_xlsx), "scheme": "scheme1"})
        elif scheme == "scheme2":
            print("[QUOTE] 开始生成结构化报价（第二方案）...")
            
            # 处理税率
            tax_rate_value = 0.0
            if tax_rate == "3%":
                tax_rate_value = 0.03
            elif tax_rate == "13%":
                tax_rate_value = 0.13
            elif tax_rate == "不含税":
                tax_rate_value = 0.0
            else:
                tax_rate_value = 0.13  # 默认13%
            
            structured_file = generate_structured_quote(
                inquiry_path=standard_csv,
                price_path=price_path,
                output_dir=quote_dir,
                customer_id=username,
                customer_name=company_name or company,
                selected_brand=brand,
                discount=user_discount,
                # 传递公司信息
                company_info={
                    "company_name": company_name,
                    "business_contact": business_contact,
                    "contact_phone": contact_phone,
                    "contact_email": contact_email,
                    "customer_header": customer_header,
                    "recipient": recipient,
                    "contact_method": contact_method,
                    "address": address,
                    "tax_rate": tax_rate_value
                }
            )
            # 对第二方案应用折扣
            try:
                import pandas as pd
                df2 = pd.read_excel(structured_file)
                if '单价' in df2.columns:
                    df2['单价'] = pd.to_numeric(df2['单价'], errors='coerce') * user_discount
                    if '数量' in df2.columns:
                        qty = pd.to_numeric(df2['数量'], errors='coerce').fillna(0)
                        df2['金额'] = (df2['单价'].fillna(0) * qty).round(2)
                    df2.to_excel(structured_file, index=False)
            except Exception as _:
                pass
            print(f"🎉 [QUOTE] 结构化报价生成成功: {os.path.basename(structured_file)}")
            result_payload.update({"structured_file": os.path.basename(structured_file), "scheme": "scheme2"})
            # 仅输出所选方案：删除标准报价中间文件
            try:
                if os.path.exists(standard_xlsx):
                    os.remove(standard_xlsx)
                if os.path.exists(standard_csv):
                    os.remove(standard_csv)
            except Exception:
                pass
        else:
            # 非法scheme值，回退到第一方案
            if auto_fill_price:
                print(f"[QUOTE] 自动用价格表再次填充价格和品牌...")
                match_quote_with_price_table(standard_xlsx, price_path, brand)
            try:
                import pandas as pd
                df1 = pd.read_excel(standard_xlsx)
                if '单价' in df1.columns:
                    df1['单价'] = pd.to_numeric(df1['单价'], errors='coerce') * user_discount
                    if '数量' in df1.columns:
                        qty = pd.to_numeric(df1['数量'], errors='coerce').fillna(0)
                        df1['总价'] = (df1['单价'].fillna(0) * qty).round(2)
                    df1.to_excel(standard_xlsx, index=False)
            except Exception as _:
                pass
            print(f"🎉 [QUOTE] 报价单生成成功(回退到第一方案): {os.path.basename(standard_xlsx)}")
            result_payload.update({"file": os.path.basename(standard_xlsx), "scheme": "scheme1"})

        return result_payload

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [QUOTE] 生成报价单失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"生成报价单失败: {str(e)}")

@app.post("/api/generate-enhanced-quote")
async def generate_enhanced_quote(
    price_file: str = Form(...),
    inquiry_file: str = Form(...),
    username: str = Depends(verify_credentials)
):
    """使用增强价格匹配功能生成报价单"""
    try:
        print(f"🚀 [ENHANCED-API] 开始增强报价生成")
        print(f"👤 [ENHANCED-API] 用户: {username}")
        print(f"💰 [ENHANCED-API] 价格文件: {price_file}")
        print(f"📋 [ENHANCED-API] 询价文件: {inquiry_file}")
        
        user_dir = os.path.join(DATA_ROOT, username)
        price_path = os.path.join(user_dir, "价格表", price_file)
        inquiry_path = os.path.join(user_dir, "询价表", inquiry_file)
        
        # 检查文件是否存在
        if not os.path.exists(price_path):
            raise HTTPException(status_code=404, detail=f"价格表文件不存在: {price_file}")
        
        if not os.path.exists(inquiry_path):
            raise HTTPException(status_code=404, detail=f"询价表文件不存在: {inquiry_file}")
        
        # 生成输出文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_name = Path(price_file).stem  # 从价格表文件名提取公司名
        output_filename = f"{timestamp}_{company_name}_增强报价.xlsx"
        output_path = os.path.join(user_dir, "报价单", output_filename)
        
        # 使用增强的价格匹配功能
        result_file = process_quote_with_enhanced_matching(
            inquiry_file=inquiry_path,
            price_file=price_path,
            output_file=output_path,
            username=username
        )
        
        if result_file and os.path.exists(result_file):
            print(f"✅ [ENHANCED-API] 增强报价生成成功: {output_filename}")
            return {
                "message": "增强报价单生成成功",
                "file": output_filename,
                "enhanced": True
            }
        else:
            raise HTTPException(status_code=500, detail="增强报价单生成失败")
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [ENHANCED-API] 增强报价生成失败: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"增强报价生成失败: {str(e)}")

@app.post("/api/generate-multi-company-quote")
async def generate_multi_company_quote_api(
    inquiry_file: str = Form(...),
    username: str = Depends(verify_credentials)
):
    """生成多公司价格对比报价单"""
    try:
        print(f"🚀 [MULTI-API] 开始多公司报价生成")
        print(f"👤 [MULTI-API] 用户: {username}")
        print(f"📋 [MULTI-API] 询价文件: {inquiry_file}")
        
        user_dir = os.path.join(DATA_ROOT, username)
        inquiry_path = os.path.join(user_dir, "询价表", inquiry_file)
        price_dir = os.path.join(user_dir, "价格表")
        
        # 检查询价文件是否存在
        if not os.path.exists(inquiry_path):
            raise HTTPException(status_code=404, detail=f"询价表文件不存在: {inquiry_file}")
        
        # 扫描所有价格表文件
        if not os.path.exists(price_dir):
            raise HTTPException(status_code=404, detail="价格表目录不存在")
        
        price_files = {}
        supported_formats = ['.xlsx', '.xls', '.csv']
        
        for filename in os.listdir(price_dir):
            if any(filename.lower().endswith(fmt) for fmt in supported_formats):
                company_name = Path(filename).stem  # 使用文件名作为公司名
                price_files[company_name] = os.path.join(price_dir, filename)
        
        if not price_files:
            raise HTTPException(status_code=404, detail="未找到任何价格表文件")
        
        print(f"📊 [MULTI-API] 发现价格表: {list(price_files.keys())}")
        
        # 生成输出文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{timestamp}_多公司价格对比.xlsx"
        output_path = os.path.join(user_dir, "报价单", output_filename)
        
        # 使用多公司报价功能
        result_file = generate_multi_brand_quote(
            inquiry_file=inquiry_path,
            price_files=price_files,
            output_file=output_path,
            username=username
        )
        
        if result_file and os.path.exists(result_file):
            print(f"✅ [MULTI-API] 多公司报价生成成功: {output_filename}")
            return {
                "message": "多公司价格对比报价单生成成功",
                "file": output_filename,
                "companies": list(price_files.keys()),
                "company_count": len(price_files)
            }
        else:
            raise HTTPException(status_code=500, detail="多公司报价单生成失败")
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [MULTI-API] 多公司报价生成失败: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"多公司报价生成失败: {str(e)}")

@app.get("/api/download/{file_type}/{filename}")
async def download_file(file_type: str, filename: str, username: str = Depends(verify_credentials)):
    """下载文件"""
    print(f"🔽 下载请求: file_type={file_type}, filename={filename}, username={username}")
    
    if file_type not in ["价格表", "询价表", "报价单"]:
        print(f"❌ 无效的文件类型: {file_type}")
        raise HTTPException(status_code=400, detail="无效的文件类型")
    
    file_path = os.path.join(DATA_ROOT, username, file_type, filename)
    print(f"📁 文件路径: {file_path}")
    print(f"📁 文件存在: {os.path.exists(file_path)}")
    
    # 列出目录内容进行调试
    dir_path = os.path.join(DATA_ROOT, username, file_type)
    if os.path.exists(dir_path):
        files_in_dir = os.listdir(dir_path)
        print(f"📂 目录 {dir_path} 中的文件: {files_in_dir}")
    else:
        print(f"❌ 目录不存在: {dir_path}")
    
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        raise HTTPException(status_code=404, detail=f"文件不存在: {filename}")
    
    print(f"✅ 开始下载文件: {file_path}")
    return FileResponse(file_path, filename=filename)

@app.post("/api/logout")
async def logout(username: str = Depends(verify_credentials)):
    """退出登录，清理用户的临时文件夹"""
    try:
        print(f"👋 [LOGOUT] 用户退出登录: {username}")
        
        # 清理用户的临时文件夹
        user_dir = os.path.join(DATA_ROOT, username)
        if os.path.exists(user_dir):
            # 查找并删除所有临时文件夹
            temp_dirs = []
            for item in os.listdir(user_dir):
                item_path = os.path.join(user_dir, item)
                if os.path.isdir(item_path) and (item.startswith("temp_") or "temp_" in item):
                    temp_dirs.append(item_path)
                    print(f"🔍 [LOGOUT] 发现临时目录: {item_path}")
            
            # 删除临时文件夹
            deleted_count = 0
            for temp_dir in temp_dirs:
                # 最多尝试3次
                for attempt in range(3):
                    try:
                        print(f"🗑️  [LOGOUT] 尝试删除临时目录 (尝试 {attempt+1}/3): {temp_dir}")
                        
                        # 先尝试清空目录内容，再删除目录本身
                        if os.path.exists(temp_dir):
                            # 列出目录内容
                            print(f"📂 [LOGOUT] 目录内容: {os.listdir(temp_dir) if os.path.exists(temp_dir) else '目录不存在'}")
                            
                            # 尝试强制删除
                            if sys.platform == 'win32':
                                # Windows平台
                                os.system(f'rd /s /q "{temp_dir}"')
                            else:
                                # Unix/Linux/MacOS平台
                                os.system(f'rm -rf "{temp_dir}"')
                            
                            # 检查是否删除成功
                            if not os.path.exists(temp_dir):
                                print(f"✅ [LOGOUT] 成功删除临时目录: {temp_dir}")
                                deleted_count += 1
                                break
                            else:
                                print(f"⚠️  [LOGOUT] 系统命令删除失败，尝试使用shutil")
                                shutil.rmtree(temp_dir, ignore_errors=True)
                                
                                if not os.path.exists(temp_dir):
                                    print(f"✅ [LOGOUT] 成功删除临时目录: {temp_dir}")
                                    deleted_count += 1
                                    break
                        else:
                            print(f"⚠️  [LOGOUT] 跳过目录(不存在): {temp_dir}")
                            break
                    except Exception as e:
                        print(f"⚠️  [LOGOUT] 删除临时目录失败 (尝试 {attempt+1}/3): {temp_dir}, 错误: {e}")
                        import traceback
                        traceback.print_exc()
                        
                        # 最后一次尝试失败
                        if attempt == 2:
                            print(f"❌ [LOGOUT] 删除临时目录失败，已尝试最大次数: {temp_dir}")
                        else:
                            # 等待一小段时间再重试
                            time.sleep(0.5)
            
            print(f"📊 [LOGOUT] 清理完成，共删除 {deleted_count}/{len(temp_dirs)} 个临时目录")
        else:
            print(f"⚠️  [LOGOUT] 用户目录不存在: {user_dir}")
        
        return {"message": "退出成功", "cleaned_dirs": deleted_count if 'deleted_count' in locals() else 0}
    except Exception as e:
        print(f"❌ [LOGOUT] 退出登录失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"退出登录失败: {str(e)}")

@app.post("/api/admin/cleanup-temp")
async def admin_cleanup_temp(username: str = Depends(verify_credentials)):
    """管理员功能：清理所有用户的临时文件夹"""
    # 验证管理员权限
    if not is_admin(username):
        raise HTTPException(status_code=403, detail="需要管理员权限")
    
    try:
        print(f"🧹 [ADMIN] 管理员 {username} 请求清理所有临时文件夹")
        
        # 遍历所有用户目录
        total_dirs = 0
        deleted_dirs = 0
        failed_dirs = 0
        user_stats = {}
        
        if os.path.exists(DATA_ROOT):
            for user_folder in os.listdir(DATA_ROOT):
                user_dir = os.path.join(DATA_ROOT, user_folder)
                if os.path.isdir(user_dir):
                    user_stats[user_folder] = {"found": 0, "deleted": 0, "failed": 0}
                    
                    # 查找临时文件夹
                    for item in os.listdir(user_dir):
                        item_path = os.path.join(user_dir, item)
                        if os.path.isdir(item_path) and ("temp_" in item):
                            total_dirs += 1
                            user_stats[user_folder]["found"] += 1
                            print(f"🔍 [ADMIN] 发现临时目录: {item_path}")
                            
                            # 尝试删除，最多3次
                            deleted = False
                            for attempt in range(3):
                                try:
                                    print(f"🗑️  [ADMIN] 尝试删除临时目录 (尝试 {attempt+1}/3): {item_path}")
                                    
                                    # 列出目录内容
                                    print(f"📂 [ADMIN] 目录内容: {os.listdir(item_path) if os.path.exists(item_path) else '目录不存在'}")
                                    
                                    # 尝试强制删除
                                    if sys.platform == 'win32':
                                        # Windows平台
                                        os.system(f'rd /s /q "{item_path}"')
                                    else:
                                        # Unix/Linux/MacOS平台
                                        os.system(f'rm -rf "{item_path}"')
                                    
                                    # 检查是否删除成功
                                    if not os.path.exists(item_path):
                                        print(f"✅ [ADMIN] 成功删除临时目录: {item_path}")
                                        deleted_dirs += 1
                                        user_stats[user_folder]["deleted"] += 1
                                        deleted = True
                                        break
                                    else:
                                        print(f"⚠️  [ADMIN] 系统命令删除失败，尝试使用shutil")
                                        shutil.rmtree(item_path, ignore_errors=True)
                                        
                                        if not os.path.exists(item_path):
                                            print(f"✅ [ADMIN] 成功删除临时目录: {item_path}")
                                            deleted_dirs += 1
                                            user_stats[user_folder]["deleted"] += 1
                                            deleted = True
                                            break
                                except Exception as e:
                                    print(f"⚠️  [ADMIN] 删除临时目录失败 (尝试 {attempt+1}/3): {item_path}, 错误: {e}")
                                    
                                    # 最后一次尝试失败
                                    if attempt == 2:
                                        print(f"❌ [ADMIN] 删除临时目录失败，已尝试最大次数: {item_path}")
                                    else:
                                        # 等待一小段时间再重试
                                        time.sleep(0.5)
                            
                            # 如果所有尝试都失败
                            if not deleted:
                                failed_dirs += 1
                                user_stats[user_folder]["failed"] += 1
                                print(f"❌ [ADMIN] 删除临时目录最终失败: {item_path}")
        
        # 生成结果报告
        result = {
            "message": f"清理完成，共发现 {total_dirs} 个临时目录，成功删除 {deleted_dirs} 个，失败 {failed_dirs} 个",
            "total_found": total_dirs,
            "total_deleted": deleted_dirs,
            "total_failed": failed_dirs,
            "user_stats": user_stats
        }
        
        print(f"📊 [ADMIN] {result['message']}")
        return result
    
    except Exception as e:
        print(f"❌ [ADMIN] 清理临时文件夹失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"清理临时文件夹失败: {str(e)}")

def cleanup_temp_directories():
    """清理所有临时文件夹"""
    try:
        print(f"🧹 [SYSTEM] 系统启动时清理临时文件夹")
        
        # 遍历所有用户目录
        total_dirs = 0
        deleted_dirs = 0
        failed_dirs = 0
        
        if os.path.exists(DATA_ROOT):
            for user_folder in os.listdir(DATA_ROOT):
                user_dir = os.path.join(DATA_ROOT, user_folder)
                if os.path.isdir(user_dir):
                    # 查找临时文件夹
                    for item in os.listdir(user_dir):
                        item_path = os.path.join(user_dir, item)
                        if os.path.isdir(item_path) and ("temp_" in item):
                            total_dirs += 1
                            print(f"🔍 [SYSTEM] 发现临时目录: {item_path}")
                            
                            # 尝试删除
                            try:
                                print(f"🗑️  [SYSTEM] 尝试删除临时目录: {item_path}")
                                
                                # 尝试强制删除
                                if sys.platform == 'win32':
                                    # Windows平台
                                    os.system(f'rd /s /q "{item_path}"')
                                else:
                                    # Unix/Linux/MacOS平台
                                    os.system(f'rm -rf "{item_path}"')
                                
                                # 检查是否删除成功
                                if not os.path.exists(item_path):
                                    print(f"✅ [SYSTEM] 成功删除临时目录: {item_path}")
                                    deleted_dirs += 1
                                else:
                                    print(f"⚠️  [SYSTEM] 系统命令删除失败，尝试使用shutil")
                                    shutil.rmtree(item_path, ignore_errors=True)
                                    
                                    if not os.path.exists(item_path):
                                        print(f"✅ [SYSTEM] 成功删除临时目录: {item_path}")
                                        deleted_dirs += 1
                                    else:
                                        print(f"❌ [SYSTEM] 删除临时目录失败: {item_path}")
                                        failed_dirs += 1
                            except Exception as e:
                                print(f"❌ [SYSTEM] 删除临时目录失败: {item_path}, 错误: {e}")
                                failed_dirs += 1
        
        print(f"📊 [SYSTEM] 清理完成，共发现 {total_dirs} 个临时目录，成功删除 {deleted_dirs} 个，失败 {failed_dirs} 个")
        return deleted_dirs
    
    except Exception as e:
        print(f"❌ [SYSTEM] 清理临时文件夹失败: {e}")
        import traceback
        traceback.print_exc()
        return 0

def match_quote_with_price_table(quote_file, price_file, selected_brand=None):
    """
    用已生成的报价单与价格表再次匹配，直接在报价单的价格列填入匹配到的价格，未匹配的保持空白，品牌列输出所选品牌。
    匹配时支持模糊包含、去除空格、全小写，遍历所选品牌的价格表行，找到最优匹配。
    忽略询价表品牌，全部用 selected_brand。
    """
    import pandas as pd
    import re
    quote_df = pd.read_excel(quote_file)
    price_df = pd.read_excel(price_file)

    # 自动识别标准型号、名称、品牌列
    model_col = None
    for col in ['标准型号', '型号', '产品型号', '规格型号', '型号编码']:
        if col in quote_df.columns and col in price_df.columns:
            model_col = col
            break
    name_col = None
    for col in ['名称', '品名', '产品名称', '项目名称', '物料名称']:
        if col in quote_df.columns and col in price_df.columns:
            name_col = col
            break
    brand_col = None
    for col in ['品牌', '厂商', '生产厂家']:
        if col in price_df.columns:
            brand_col = col
            break

    def normalize(s):
        if pd.isna(s):
            return ''
        return re.sub(r'\s+', '', str(s)).lower()

    # 确保总价列存在
    if '总价' not in quote_df.columns:
        quote_df['总价'] = ''
        print("[DEBUG] 添加总价列")

    for idx, row in quote_df.iterrows():
        best_score = 0
        best_match = None
        q_model = normalize(row.get(model_col, '')) if model_col else ''
        q_name = normalize(row.get(name_col, '')) if name_col else ''
        q_brand = normalize(selected_brand) if selected_brand else ''  # 只用所选品牌
        # 只遍历选中品牌的行
        if selected_brand and brand_col:
            price_rows = price_df[price_df[brand_col].apply(lambda x: normalize(x) == q_brand)]
        else:
            price_rows = price_df
        for _, prow in price_rows.iterrows():
            p_model = normalize(prow.get(model_col, '')) if model_col else ''
            p_name = normalize(prow.get(name_col, '')) if name_col else ''
            p_brand = normalize(prow.get(brand_col, '')) if brand_col else ''
            score = 0
            # 完全相等优先
            if q_model and p_model and q_model == p_model:
                score += 10
            elif q_model and p_model and (q_model in p_model or p_model in q_model):
                score += 6
            if q_name and p_name and q_name == p_name:
                score += 5
            elif q_name and p_name and (q_name in p_name or p_name in q_name):
                score += 3
            # 品牌只用所选品牌，不再考虑询价表品牌
            if q_brand and p_brand and q_brand == p_brand:
                score += 2
            elif q_brand and p_brand and (q_brand in p_brand or p_brand in q_brand):
                score += 1
            if score > best_score:
                best_score = score
                best_match = prow
        # 匹配到则写入单价和品牌，否则保持空白
        if best_match is not None and best_score > 0:
            unit_price = best_match.get('单价', '')
            quote_df.at[idx, '单价'] = unit_price
            # 计算总价 = 单价 × 数量
            try:
                quantity = row.get('数量', '')
                if quantity and str(quantity).strip() != '':
                    qty_val = float(quantity)
                    price_val = float(unit_price) if unit_price and str(unit_price).strip() != '' else 0
                    total_val = price_val * qty_val
                    quote_df.at[idx, '总价'] = total_val
                    print(f"[DEBUG] 行{idx+1}: {price_val} × {qty_val} = {total_val}")
                else:
                    quote_df.at[idx, '总价'] = 0.0
                    print(f"[DEBUG] 行{idx+1}: 数量为空，总价=0.0")
            except Exception as e:
                print(f"[警告] 总价计算失败: 单价={unit_price}, 数量={quantity}, 错误={e}")
                quote_df.at[idx, '总价'] = 0.0
            
            if brand_col:
                quote_df.at[idx, '品牌'] = selected_brand if selected_brand else best_match.get(brand_col, '')
        else:
            # 没匹配到也要把品牌列写成所选品牌
            if brand_col and selected_brand:
                quote_df.at[idx, '品牌'] = selected_brand
            # 没匹配到也要计算总价（如果有单价和数量）
            try:
                unit_price = row.get('单价', '')
                quantity = row.get('数量', '')
                if unit_price and quantity and str(unit_price).strip() != '' and str(quantity).strip() != '':
                    price_val = float(unit_price)
                    qty_val = float(quantity)
                    total_val = price_val * qty_val
                    quote_df.at[idx, '总价'] = total_val
                    print(f"[DEBUG] 未匹配行{idx+1}: {price_val} × {qty_val} = {total_val}")
                else:
                    quote_df.at[idx, '总价'] = 0.0
                    print(f"[DEBUG] 未匹配行{idx+1}: 单价或数量为空，总价=0.0")
            except Exception as e:
                print(f"[警告] 未匹配行{idx+1}: 总价计算失败 - {e}")
                quote_df.at[idx, '总价'] = 0.0
    
    # 最终验证和重新计算总价
    print("[DEBUG] 开始最终总价验证和重新计算")
    for idx, row in quote_df.iterrows():
        try:
            # 只使用单价列计算总价
            unit_price = row.get('单价', '')
            quantity = row.get('数量', '')
            
            # 计算总价 = 单价 × 数量
            if unit_price and quantity and str(unit_price).strip() != '' and str(quantity).strip() != '':
                try:
                    price_val = float(unit_price)
                    qty_val = float(quantity)
                    total_val = price_val * qty_val
                    quote_df.at[idx, '总价'] = total_val
                    print(f"[DEBUG] 最终计算: 行{idx+1}, {price_val} × {qty_val} = {total_val}")
                except ValueError as e:
                    print(f"[警告] 最终计算失败: 行{idx+1}, 数值转换失败 - {e}")
                    quote_df.at[idx, '总价'] = 0.0
            else:
                print(f"[警告] 最终计算失败: 行{idx+1}, 单价或数量为空")
                quote_df.at[idx, '总价'] = 0.0
                
        except Exception as e:
            print(f"[错误] 最终计算失败: 行{idx+1}, 计算失败 - {e}")
            quote_df.at[idx, '总价'] = 0.0
    
    quote_df.to_excel(quote_file, index=False)
    return quote_file

@app.get("/api/price-table/{filename}")
async def get_price_table_content(filename: str, username: str = Depends(verify_credentials)):
    """获取价格表内容"""
    try:
        # 构建文件路径
        user_dir = os.path.join(DATA_ROOT, username, "价格表")
        file_path = os.path.join(user_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取Excel文件
        df = pd.read_excel(file_path)
        
        # 处理NaN值，将其替换为None，以便JSON序列化
        df_cleaned = df.where(pd.notna(df), None)
        
        # 转换为字典格式
        data = {
            "columns": df_cleaned.columns.tolist(),
            "data": df_cleaned.values.tolist(),
            "total_rows": len(df_cleaned),
            "total_columns": len(df_cleaned.columns)
        }
        
        return JSONResponse(content=data)
        
    except Exception as e:
        print(f"❌ [ERROR] 获取价格表内容失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取价格表内容失败: {str(e)}")

@app.post("/api/price-table/{filename}/update")
async def update_price_table_content(
    filename: str, 
    data: Dict[str, Any], 
    username: str = Depends(verify_credentials)
):
    """更新价格表内容"""
    try:
        # 构建文件路径
        user_dir = os.path.join(DATA_ROOT, username, "价格表")
        file_path = os.path.join(user_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 从请求数据中提取列名和数据
        columns = data.get("columns", [])
        rows = data.get("data", [])
        
        # 创建DataFrame
        df = pd.DataFrame(rows, columns=columns)
        
        # 保存到文件
        df.to_excel(file_path, index=False)
        
        return JSONResponse(content={"message": "价格表更新成功"})
        
    except Exception as e:
        print(f"❌ [ERROR] 更新价格表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"更新价格表失败: {str(e)}")

@app.delete("/api/price-table/{filename}")
async def delete_price_table(filename: str, username: str = Depends(verify_credentials)):
    """删除价格表"""
    try:
        # 构建文件路径
        user_dir = os.path.join(DATA_ROOT, username, "价格表")
        file_path = os.path.join(user_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 删除文件
        os.remove(file_path)
        
        return JSONResponse(content={"message": "价格表删除成功"})
        
    except Exception as e:
        print(f"❌ [ERROR] 删除价格表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除价格表失败: {str(e)}")

@app.post("/api/price-table/{filename}/add-row")
async def add_row_to_price_table(
    filename: str, 
    row_data: Dict[str, Any], 
    username: str = Depends(verify_credentials)
):
    """向价格表添加新行"""
    try:
        # 构建文件路径
        user_dir = os.path.join(DATA_ROOT, username, "价格表")
        file_path = os.path.join(user_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取现有数据
        df = pd.read_excel(file_path)
        
        # 添加新行
        new_row = pd.DataFrame([row_data])
        df = pd.concat([df, new_row], ignore_index=True)
        
        # 保存文件
        df.to_excel(file_path, index=False)
        
        return JSONResponse(content={"message": "行添加成功", "new_row_index": len(df) - 1})
        
    except Exception as e:
        print(f"❌ [ERROR] 添加行失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"添加行失败: {str(e)}")

@app.put("/api/price-table/{filename}/row/{row_index}")
async def update_row_in_price_table(
    filename: str, 
    row_index: int, 
    row_data: Dict[str, Any], 
    username: str = Depends(verify_credentials)
):
    """更新价格表中的指定行"""
    try:
        # 构建文件路径
        user_dir = os.path.join(DATA_ROOT, username, "价格表")
        file_path = os.path.join(user_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取现有数据
        df = pd.read_excel(file_path)
        
        if row_index >= len(df):
            raise HTTPException(status_code=400, detail="行索引超出范围")
        
        # 更新指定行
        for column, value in row_data.items():
            if column in df.columns:
                df.at[row_index, column] = value
        
        # 保存文件
        df.to_excel(file_path, index=False)
        
        return JSONResponse(content={"message": "行更新成功"})
        
    except Exception as e:
        print(f"❌ [ERROR] 更新行失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"更新行失败: {str(e)}")

@app.delete("/api/price-table/{filename}/row/{row_index}")
async def delete_row_from_price_table(
    filename: str, 
    row_index: int, 
    username: str = Depends(verify_credentials)
):
    """删除价格表中的指定行"""
    try:
        # 构建文件路径
        user_dir = os.path.join(DATA_ROOT, username, "价格表")
        file_path = os.path.join(user_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取现有数据
        df = pd.read_excel(file_path)
        
        if row_index >= len(df):
            raise HTTPException(status_code=400, detail="行索引超出范围")
        
        # 删除指定行
        df = df.drop(row_index).reset_index(drop=True)
        
        # 保存文件
        df.to_excel(file_path, index=False)
        
        return JSONResponse(content={"message": "行删除成功"})
        
    except Exception as e:
        print(f"❌ [ERROR] 删除行失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除行失败: {str(e)}")

@app.post("/api/price-table/{filename}/add-column")
async def add_column_to_price_table(
    filename: str, 
    data: Dict[str, Any],
    username: str = Depends(verify_credentials)
):
    """向价格表添加新列"""
    try:
        # 构建文件路径
        user_dir = os.path.join(DATA_ROOT, username, "价格表")
        file_path = os.path.join(user_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取现有数据
        df = pd.read_excel(file_path)
        
        # 从请求数据中获取列名
        column_name = data.get("column_name", "")
        if not column_name:
            raise HTTPException(status_code=400, detail="列名不能为空")
        
        # 检查列名是否已存在
        if column_name in df.columns:
            raise HTTPException(status_code=400, detail="列名已存在")
        
        # 添加新列
        df[column_name] = ""
        
        # 保存文件
        df.to_excel(file_path, index=False)
        
        return JSONResponse(content={"message": "列添加成功"})
        
    except Exception as e:
        print(f"❌ [ERROR] 添加列失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"添加列失败: {str(e)}")

@app.delete("/api/price-table/{filename}/column/{column_name}")
async def delete_column_from_price_table(
    filename: str, 
    column_name: str, 
    username: str = Depends(verify_credentials)
):
    """删除价格表中的指定列"""
    try:
        # 构建文件路径
        user_dir = os.path.join(DATA_ROOT, username, "价格表")
        file_path = os.path.join(user_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="文件不存在")
        
        # 读取现有数据
        df = pd.read_excel(file_path)
        
        if column_name not in df.columns:
            raise HTTPException(status_code=400, detail="列名不存在")
        
        # 删除指定列
        df = df.drop(columns=[column_name])
        
        # 保存文件
        df.to_excel(file_path, index=False)
        
        return JSONResponse(content={"message": "列删除成功"})
        
    except Exception as e:
        print(f"❌ [ERROR] 删除列失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"删除列失败: {str(e)}")

@app.post("/api/generate-structured-quote")
async def generate_structured_quote_api(
    price_file: str = Form(...),
    inquiry_file: str = Form(...),
    customer_id: str = Form(...),
    customer_name: str = Form(None),
    currency: str = Form("CNY"),
    tax_rate: float = Form(0.13),
    valid_days: int = Form(30),
    lead_time: str = Form("7-10 天"),
    payment_terms: str = Form("款到发货"),
    sales_contact: str = Form(None),
    sales_phone: str = Form(None),
    sales_email: str = Form(None),
    username: str = Depends(verify_credentials)
):
    try:
        user_dir = os.path.join(DATA_ROOT, username)
        price_path = os.path.join(user_dir, "价格表", price_file)
        inquiry_path = os.path.join(user_dir, "询价表", inquiry_file)
        output_dir = os.path.join(user_dir, "报价单")
        os.makedirs(output_dir, exist_ok=True)

        if not os.path.exists(price_path):
            raise HTTPException(status_code=404, detail=f"价格文件不存在: {price_file}")
        if not os.path.exists(inquiry_path):
            raise HTTPException(status_code=404, detail=f"询价文件不存在: {inquiry_file}")

        output_file = generate_structured_quote(
            inquiry_path=inquiry_path,
            price_path=price_path,
            output_dir=output_dir,
            customer_id=customer_id,
            customer_name=customer_name,
            currency=currency,
            tax_rate=tax_rate,
            valid_days=valid_days,
            lead_time=lead_time,
            payment_terms=payment_terms,
            sales_contact=sales_contact,
            sales_phone=sales_phone,
            sales_email=sales_email,
            discount=user_discount,
        )

        filename = os.path.basename(output_file)
        return {"message": "结构化报价单生成成功", "file": filename}
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [STRUCTURED] 生成结构化报价失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"生成结构化报价失败: {str(e)}")

# -----------------------------
# 管理员设置用户折扣
# -----------------------------
@app.get("/api/admin/user-discount")
async def admin_get_user_discount(target_user: str = Query(...), username: str = Depends(verify_credentials)):
    if not is_admin(username):
        raise HTTPException(status_code=403, detail="需要管理员权限")
    rules_manager = get_rules_manager()
    discount = rules_manager.get_user_discount(target_user)
    return {"user": target_user, "discount": discount}

@app.post("/api/admin/user-discount")
async def admin_set_user_discount(target_user: str = Form(...), discount: float = Form(...), username: str = Depends(verify_credentials)):
    if not is_admin(username):
        raise HTTPException(status_code=403, detail="需要管理员权限")
    rules_manager = get_rules_manager()
    success = rules_manager.set_user_discount(target_user, discount)
    if not success:
        raise HTTPException(status_code=400, detail="设置折扣失败")
    return {"user": target_user, "discount": rules_manager.get_user_discount(target_user)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)